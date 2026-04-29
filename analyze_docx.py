"""
Analyze all 11 Dify output docx files from the 奇瑞HODE01 output directory.
Extract structured summary and sample test cases.
Writes results to a temp file.
"""

import os
import sys
from docx import Document

BASE_DIR = r"F:\UDS_ASSISTANT\output\奇瑞HODE01"
SERVICES = ["0X10","0X11","0X14","0X19","0X22","0X27","0X28","0X2E","0X31","0X3E","0X85"]
OUTPUT_FILE = r"F:\UDS_ASSISTANT\analysis_output.txt"

def extract_paragraphs(docx_path):
    """Return all paragraphs as list of (style_name, text)."""
    doc = Document(docx_path)
    result = []
    for p in doc.paragraphs:
        result.append((p.style.name if p.style else "None", p.text))
    # Also grab tables
    tables = doc.tables
    return result, tables

def analyze_service(service_id):
    """Analyze one service's test.docx and return structured info."""
    path = os.path.join(BASE_DIR, service_id, "test.docx")
    if not os.path.exists(path):
        return None, f"FILE NOT FOUND: {path}"

    try:
        doc = Document(path)
    except Exception as e:
        return None, f"ERROR opening: {e}"

    issues = []
    paragraphs = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else "None"
        text = p.text.strip()
        paragraphs.append((style, text))

    # Detect categories: ### headers => style 'Heading 3' or text starting with '###'
    categories = []
    current_category = None
    case_ids = {}  # category -> list of case IDs
    all_case_ids = []

    for style, text in paragraphs:
        # Detect heading 3 or lines starting with ###
        is_h3 = (style == "Heading 3") or (text.startswith("### ") and style in ("Normal", "None", ""))
        if is_h3:
            cat_name = text.lstrip("#").strip()
            if cat_name:
                current_category = cat_name
                categories.append(cat_name)
                case_ids[cat_name] = []

        # Detect case IDs - common patterns like "TC_XXXX", "TestCase", "Case ID:", numbered patterns
        # Also look for lines that look like case identifiers
        if text and current_category:
            # Common case ID patterns
            if any(kw in text.upper() for kw in ["CASE ID", "CASEID", "TEST CASE", "TC_", "TC-"]):
                case_ids[current_category].append(text)
                all_case_ids.append(text)
            elif text.startswith("ID:") or text.startswith("ID："):
                case_ids[current_category].append(text)
                all_case_ids.append(text)

    # If no structured case IDs found, try counting rows in tables
    table_count = len(doc.tables)
    table_rows_info = []
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        table_rows_info.append((i, rows, cols))

    # Check for encoding issues
    full_text = "\n".join(t for _, t in paragraphs)
    encoding_issues = []
    for bad in ["�", "???", "ï¿½"]:
        if bad in full_text:
            encoding_issues.append(f"Found replacement character: {repr(bad)}")

    # Check for empty sections
    empty_lines = sum(1 for _, t in paragraphs if not t)

    # Gather heading styles used
    heading_styles = set()
    for style, text in paragraphs:
        if "Heading" in style or text.startswith("#"):
            heading_styles.add(style)

    return {
        "service": service_id,
        "path": path,
        "total_paragraphs": len(paragraphs),
        "empty_paragraphs": empty_lines,
        "num_tables": table_count,
        "tables_info": table_rows_info,
        "categories": categories,
        "case_ids": case_ids,
        "all_case_ids": all_case_ids,
        "heading_styles": heading_styles,
        "encoding_issues": encoding_issues,
        "full_text": full_text,
        "paragraphs": paragraphs,
    }, None

def extract_first_n_cases(service_id, n=3):
    """Extract first N complete test cases with steps and expected output."""
    path = os.path.join(BASE_DIR, service_id, "test.docx")
    doc = Document(path)

    paragraphs = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else "None"
        text = p.text.strip()
        paragraphs.append((style, text))

    # Find case blocks - look for patterns that delimit test cases
    # Cases are usually separated by headings or specific markers
    cases = []
    current_case = None
    current_lines = []

    # Strategy: split by lines that look like case IDs or headings
    for style, text in paragraphs:
        is_case_start = False

        # Check if this line starts a new test case
        if any(kw in text.upper() for kw in ["CASE ID", "CASEID", "TEST CASE"]) and text.strip():
            is_case_start = True
        elif text.startswith("ID:") or text.startswith("ID："):
            is_case_start = True
        elif style == "Heading 3" or (text.startswith("###") and text.strip()):
            # Category headers can also start new sections
            is_case_start = True

        if is_case_start and current_lines:
            # Save previous case
            cases.append(current_lines)
            current_lines = [(style, text)]
        elif text.strip():
            if current_lines:
                current_lines.append((style, text))
            else:
                current_lines = [(style, text)]

    if current_lines:
        cases.append(current_lines)

    return cases[:n]

def extract_cases_from_tables(service_id, n=3):
    """Try to extract test cases from tables (common format in docx)."""
    path = os.path.join(BASE_DIR, service_id, "test.docx")
    doc = Document(path)

    results = []
    for ti, table in enumerate(doc.tables):
        if len(results) >= n:
            break
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)
        if rows_data:
            results.append({
                "table_index": ti,
                "rows": rows_data
            })
    return results

# ==================== MAIN ====================
with open(OUTPUT_FILE, "w", encoding="utf-8") as out:
    out.write("=" * 100 + "\n")
    out.write("ANALYSIS OF DIFY OUTPUT DOCX FILES - 奇瑞HODE01\n")
    out.write("=" * 100 + "\n\n")

    all_results = {}

    # --- Part 1: Summary table per service ---
    out.write("PART 1: SUMMARY TABLE\n")
    out.write("-" * 100 + "\n\n")

    for svc in SERVICES:
        info, err = analyze_service(svc)
        if err:
            out.write(f"Service {svc}: {err}\n\n")
            continue

        all_results[svc] = info
        out.write(f"SERVICE: {svc}\n")
        out.write(f"  File: {info['path']}\n")
        out.write(f"  Total paragraphs: {info['total_paragraphs']}\n")
        out.write(f"  Empty paragraphs: {info['empty_paragraphs']}\n")
        out.write(f"  Number of tables: {info['num_tables']}\n")
        for ti, rows, cols in info['tables_info']:
            out.write(f"    Table {ti}: {rows} rows x {cols} cols\n")
        out.write(f"  Heading styles used: {', '.join(sorted(info['heading_styles']))}\n")
        out.write(f"  Categories ({len(info['categories'])}):\n")
        for cat in info['categories']:
            cid_count = len(info['case_ids'].get(cat, []))
            sample_ids = info['case_ids'].get(cat, [])[:3]
            out.write(f"    - {cat}: {cid_count} explicit case IDs")
            if sample_ids:
                out.write(f" (samples: {'; '.join(sample_ids)})")
            out.write("\n")

        total_explicit = len(info['all_case_ids'])
        out.write(f"  Total explicit case IDs found: {total_explicit}\n")

        # Estimate total cases from tables
        total_table_body_rows = sum(max(0, rows - 1) for _, rows, _ in info['tables_info'])
        if total_table_body_rows > 0:
            out.write(f"  Estimated cases from table data rows: {total_table_body_rows}\n")

        if info['encoding_issues']:
            out.write(f"  ENCODING ISSUES: {'; '.join(info['encoding_issues'])}\n")
        else:
            out.write(f"  Encoding issues: None detected\n")

        out.write("\n")

    # --- Summary comparison table ---
    out.write("\n" + "=" * 100 + "\n")
    out.write("COMPARISON TABLE\n")
    out.write("=" * 100 + "\n\n")

    header = f"{'Service':<8} {'Paras':>6} {'Tables':>6} {'TblRows':>7} {'Categories':>10} {'ExplicitIDs':>11} {'EncIssues':>10}"
    out.write(header + "\n")
    out.write("-" * len(header) + "\n")

    for svc in SERVICES:
        if svc not in all_results:
            out.write(f"{svc:<8} {'N/A':>6}\n")
            continue
        info = all_results[svc]
        total_rows = sum(r for _, r, _ in info['tables_info'])
        enc = "Yes" if info['encoding_issues'] else "None"
        out.write(f"{svc:<8} {info['total_paragraphs']:>6} {info['num_tables']:>6} {total_rows:>7} {len(info['categories']):>10} {len(info['all_case_ids']):>11} {enc:>10}\n")

    # --- Part 2: Detailed first 3 test cases for 0x10, 0x11, 0x27 ---
    out.write("\n\n" + "=" * 100 + "\n")
    out.write("PART 2: FIRST 3 TEST CASES FOR 0X10, 0X11, 0X27\n")
    out.write("=" * 100 + "\n\n")

    for svc in ["0X10", "0X11", "0X27"]:
        out.write(f"\n{'#' * 60}\n")
        out.write(f"SERVICE: {svc}\n")
        out.write(f"{'#' * 60}\n\n")

        # First try table-based extraction
        table_cases = extract_cases_from_tables(svc, n=3)
        if table_cases:
            out.write("--- TABLE-BASED EXTRACTION ---\n\n")
            for tc in table_cases:
                out.write(f"Table {tc['table_index']}:\n")
                for ri, row in enumerate(tc['rows']):
                    out.write(f"  Row {ri}: {' | '.join(row)}\n")
                out.write("\n")

        # Also try paragraph-based extraction
        out.write("--- PARAGRAPH-BASED EXTRACTION (first ~30 paragraphs) ---\n\n")
        path = os.path.join(BASE_DIR, svc, "test.docx")
        if os.path.exists(path):
            doc = Document(path)
            paras = list(doc.paragraphs)[:30]
            for i, p in enumerate(paras):
                style = p.style.name if p.style else "None"
                text = p.text
                if text.strip():
                    prefix = f"[{i}] ({style}) "
                    out.write(f"{prefix}{text}\n")
            out.write("\n")

    # --- Part 3: Raw content sample for each service (first 5 lines) ---
    out.write("\n\n" + "=" * 100 + "\n")
    out.write("PART 3: FIRST 5 NON-EMPTY LINES PER SERVICE (quick format check)\n")
    out.write("=" * 100 + "\n\n")

    for svc in SERVICES:
        if svc not in all_results:
            continue
        info = all_results[svc]
        out.write(f"SERVICE {svc}:\n")
        count = 0
        for style, text in info['paragraphs']:
            if text.strip():
                out.write(f"  ({style}) {text[:120]}\n")
                count += 1
                if count >= 5:
                    break
        out.write("\n")

print(f"Analysis written to {OUTPUT_FILE}")
